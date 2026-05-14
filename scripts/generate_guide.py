"""
Generates AI Agent Strategy Guide as:
  - /home/user/Desktop/AI_Agent_Strategy_Guide.docx
  - /home/user/Desktop/AI_Agent_Strategy_Guide.pdf
"""

import os
import io
import subprocess
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colours ──────────────────────────────────────────────────────────────────
C_NAVY   = "#0f3460"
C_RED    = "#e94560"
C_GOLD   = "#f59e0b"
C_GREEN  = "#10b981"
C_TEAL   = "#0ea5e9"
C_PURPLE = "#8b5cf6"
C_GREY   = "#94a3b8"

def hex_rgb(h):
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16)/255 for i in (0, 2, 4))

NAVY   = hex_rgb(C_NAVY)
RED    = hex_rgb(C_RED)
GOLD   = hex_rgb(C_GOLD)
GREEN  = hex_rgb(C_GREEN)
TEAL   = hex_rgb(C_TEAL)
PURPLE = hex_rgb(C_PURPLE)
GREY   = hex_rgb(C_GREY)

OUT_DIR = os.path.expanduser("~/Desktop")
os.makedirs(OUT_DIR, exist_ok=True)
DOCX_PATH = os.path.join(OUT_DIR, "AI_Agent_Strategy_Guide.docx")
CHART_DIR = "/tmp/guide_charts"
os.makedirs(CHART_DIR, exist_ok=True)

# ═════════════════════════════════════════════════════════════════════════════
# 1. CHART GENERATION
# ═════════════════════════════════════════════════════════════════════════════

def save_fig(name, fig):
    path = os.path.join(CHART_DIR, name)
    fig.savefig(path, dpi=160, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def chart_online_categories():
    fig, ax = plt.subplots(figsize=(5, 4))
    labels = ["Lead Generation\n& Traffic", "Booking &\nConversion",
              "Customer\nRetention", "Operations &\nUpselling"]
    colours = [NAVY, RED, GREEN, GOLD]
    wedges, texts, autotexts = ax.pie(
        [5, 5, 5, 5], labels=labels, colors=colours,
        autopct="%1.0f%%", startangle=90,
        textprops={"fontsize": 9}, pctdistance=0.75,
        wedgeprops={"linewidth": 2, "edgecolor": "white"})
    for at in autotexts:
        at.set_color("white"); at.set_fontsize(9); at.set_fontweight("bold")
    ax.set_title("Online AI Agents by Category\n(5 agents each)", fontsize=11, fontweight="bold", color=NAVY)
    return save_fig("online_categories.png", fig)


def chart_online_impact():
    fig, ax = plt.subplots(figsize=(6, 4))
    cats = ["Lead Gen", "Booking", "Retention", "Operations"]
    x = np.arange(len(cats)); w = 0.35
    b1 = ax.bar(x - w/2, [8, 9, 7, 6], w, label="Revenue Impact", color=[NAVY, RED, GREEN, GOLD], edgecolor="white")
    b2 = ax.bar(x + w/2, [7, 8, 9, 5], w, label="Ease of Deployment",
                color=[(*NAVY, 0.45), (*RED, 0.45), (*GREEN, 0.45), (*GOLD, 0.45)], edgecolor="white")
    ax.set_ylim(0, 11); ax.set_xticks(x); ax.set_xticklabels(cats, fontsize=9)
    ax.set_ylabel("Score (0–10)"); ax.legend(fontsize=8)
    ax.set_title("Online Agent Impact vs Ease of Deployment", fontsize=11, fontweight="bold", color=NAVY)
    ax.spines[["top","right"]].set_visible(False)
    for bar in list(b1) + list(b2):
        ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.1,
                str(int(bar.get_height())), ha="center", va="bottom", fontsize=8, fontweight="bold")
    return save_fig("online_impact.png", fig)


def chart_instore_categories():
    fig, ax = plt.subplots(figsize=(5, 4))
    labels = ["Check-In &\nReception", "In-Store Sales\n& Upselling",
              "Technician\nSupport", "Customer\nHandoff"]
    colours = [TEAL, PURPLE, RED, GOLD]
    wedges, texts, autotexts = ax.pie(
        [5, 5, 5, 5], labels=labels, colors=colours,
        autopct="%1.0f%%", startangle=90,
        textprops={"fontsize": 9}, pctdistance=0.75,
        wedgeprops={"linewidth": 2, "edgecolor": "white"})
    for at in autotexts:
        at.set_color("white"); at.set_fontsize(9); at.set_fontweight("bold")
    ax.set_title("In-Store AI Agents by Category\n(5 agents each)", fontsize=11, fontweight="bold", color=NAVY)
    return save_fig("instore_categories.png", fig)


def chart_instore_top3():
    fig, ax = plt.subplots(figsize=(6, 4))
    agents = ["#1 Smart\nCheck-In Kiosk", "#11 AI Repair\nGuide", "#17 On-the-Spot\nReview"]
    x = np.arange(len(agents)); w = 0.25
    ax.bar(x - w, [9, 8, 5], w, label="Staff Time Saved", color=NAVY, edgecolor="white")
    ax.bar(x,     [6, 7, 9], w, label="Revenue / Review Impact", color=RED, edgecolor="white")
    ax.bar(x + w, [7, 8, 9], w, label="Ease of Deployment", color=GREEN, edgecolor="white")
    ax.set_ylim(0, 11); ax.set_xticks(x); ax.set_xticklabels(agents, fontsize=9)
    ax.set_ylabel("Score (0–10)"); ax.legend(fontsize=8)
    ax.set_title("Top 3 In-Store Agents — Priority Scores", fontsize=11, fontweight="bold", color=NAVY)
    ax.spines[["top","right"]].set_visible(False)
    return save_fig("instore_top3.png", fig)


def chart_weekly_sales():
    fig, ax = plt.subplots(figsize=(5, 4))
    labels = ["Worst Week\n(Mar 30–Apr 5)", "Average", "Best Week\n(Mar 2–8)"]
    vals = [1818, 2314, 2834]
    colours = [RED, GOLD, GREEN]
    bars = ax.bar(labels, vals, color=colours, edgecolor="white", width=0.55)
    ax.set_ylim(1500, 3100)
    ax.set_ylabel("Weekly Sales (£)")
    ax.set_title("Weekly Sales Performance", fontsize=11, fontweight="bold", color=NAVY)
    ax.spines[["top","right"]].set_visible(False)
    for bar, val in zip(bars, vals):
        ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+20,
                f"£{val:,}", ha="center", va="bottom", fontsize=10, fontweight="bold")
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda v, _: f"£{int(v):,}"))
    return save_fig("weekly_sales.png", fig)


def chart_payment_split():
    fig, ax = plt.subplots(figsize=(4.5, 4))
    wedges, texts, autotexts = ax.pie(
        [59, 41], labels=["Cash", "Card"],
        colors=[NAVY, TEAL], autopct="%1.0f%%",
        startangle=90, pctdistance=0.6,
        wedgeprops={"linewidth": 3, "edgecolor": "white"},
        textprops={"fontsize": 12, "fontweight": "bold"})
    for at in autotexts:
        at.set_color("white"); at.set_fontsize(13); at.set_fontweight("bold")
    ax.set_title("Payment Method Split", fontsize=11, fontweight="bold", color=NAVY)
    return save_fig("payment_split.png", fig)


def chart_cost_breakdown():
    fig, ax = plt.subplots(figsize=(5.5, 4.5))
    labels = ["Salaries\n(£800/wk)", "Bangla Tech\nParts", "Revive\nParts",
              "eBay\nParts", "HMRC / Tax", "Council\nRates", "Insurance\n& Other"]
    vals = [35, 18, 12, 10, 8, 7, 10]
    colours = [RED, NAVY, TEAL, PURPLE, GOLD, GREEN, GREY]
    wedges, texts, autotexts = ax.pie(
        vals, labels=labels, colors=colours,
        autopct="%1.0f%%", startangle=140,
        textprops={"fontsize": 8}, pctdistance=0.78,
        wedgeprops={"linewidth": 2, "edgecolor": "white"})
    for at in autotexts:
        at.set_color("white"); at.set_fontsize(8); at.set_fontweight("bold")
    ax.set_title("Weekly Cost Breakdown\n(% of Revenue)", fontsize=11, fontweight="bold", color=NAVY)
    return save_fig("cost_breakdown.png", fig)


def chart_revenue_target():
    fig, ax = plt.subplots(figsize=(5, 4))
    labels = ["Current\nAverage", "Interim\nTarget", "Expansion\nTarget"]
    vals = [2314, 2900, 3500]
    colours = [GOLD, TEAL, GREEN]
    bars = ax.bar(labels, vals, color=colours, edgecolor="white", width=0.5)
    ax.set_ylim(2000, 4000)
    ax.set_ylabel("Weekly Revenue (£)")
    ax.set_title("Revenue vs Expansion Target (Weekly £)", fontsize=11, fontweight="bold", color=NAVY)
    ax.spines[["top","right"]].set_visible(False)
    for bar, val in zip(bars, vals):
        ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+30,
                f"£{val:,}", ha="center", va="bottom", fontsize=10, fontweight="bold")
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda v, _: f"£{int(v):,}"))
    return save_fig("revenue_target.png", fig)


def chart_revenue_strategy():
    fig, ax = plt.subplots(figsize=(7, 5))
    strategies = [
        "Part-Time Staffing Saving",
        "eBay Sales Operation",
        "Extended Weekend Hours",
        "Refurbished Phone Retail",
        "Repair Membership (50 members)",
        "Increased Avg Order Value",
        "WhatsApp Broadcast List",
    ]
    mins = [800, 400, 400, 400, 300, 200, 100]
    maxs = [1000, 1000, 800, 800, 300, 600, 400]
    y = np.arange(len(strategies))
    ax.barh(y, maxs, color=(*GREEN, 0.7), edgecolor="white", label="Max (£/month)")
    ax.barh(y, mins, color=NAVY, edgecolor="white", label="Min (£/month)")
    ax.set_yticks(y); ax.set_yticklabels(strategies, fontsize=9)
    ax.set_xlabel("Additional Monthly Revenue (£)")
    ax.set_title("Estimated Additional Monthly Revenue\nby Scaling Strategy", fontsize=11, fontweight="bold", color=NAVY)
    ax.legend(fontsize=9); ax.spines[["top","right"]].set_visible(False)
    ax.xaxis.set_major_formatter(plt.FuncFormatter(lambda v, _: f"£{int(v):,}"))
    return save_fig("revenue_strategy.png", fig)


def chart_priority_matrix():
    fig, ax = plt.subplots(figsize=(6, 5))
    items = [
        ("Fix Cash Flow",          2,  10, 200, RED),
        ("Google Business Profile",2,   9, 160, GREEN),
        ("Separate Finances",      1,   9, 130, GREEN),
        ("WhatsApp Broadcast",     1,   8, 130, TEAL),
        ("Part-Time Staffing",     4,   8, 150, TEAL),
        ("eBay Store",             4,   7, 150, GOLD),
        ("Repair Membership",      5,   7, 130, PURPLE),
        ("Refurb Retail",          6,   6, 150, NAVY),
        ("Extend Hours",           2,   6, 110, GREY),
        ("Loyalty Card",           1,   5, 100, RED),
    ]
    for label, x, y, size, colour in items:
        ax.scatter(x, y, s=size, color=colour, alpha=0.8, zorder=3)
        ax.annotate(label, (x, y), textcoords="offset points",
                    xytext=(8, 4), fontsize=7.5, color="#333")
    ax.axhline(7.5, color=GREY, linestyle="--", linewidth=0.8, alpha=0.6)
    ax.axvline(3.5, color=GREY, linestyle="--", linewidth=0.8, alpha=0.6)
    ax.text(1.2, 10.3, "Quick Wins", fontsize=8, color=GREEN, fontweight="bold")
    ax.text(5.0, 10.3, "High Effort\nHigh Reward", fontsize=8, color=GOLD, fontweight="bold")
    ax.set_xlim(0, 9); ax.set_ylim(3, 11)
    ax.set_xlabel("Effort Required  (1 = Low)", fontsize=9)
    ax.set_ylabel("Revenue Impact  (10 = High)", fontsize=9)
    ax.set_title("Strategy Priority Matrix\n(Impact vs Effort)", fontsize=11, fontweight="bold", color=NAVY)
    ax.spines[["top","right"]].set_visible(False)
    return save_fig("priority_matrix.png", fig)


def chart_roadmap():
    fig, ax = plt.subplots(figsize=(8, 3.5))
    phases = ["Phase 1\nMonth 1", "Phase 2\nMonths 2–3", "Phase 3\nMonths 3–6", "Phase 4\nMonth 6+"]
    tasks = [
        "Fix Cash Flow + Separate Finances",
        "Google Business + WhatsApp Broadcast",
        "AI Booking Bot + Status Updater + eBay Store",
        "Membership + Refurb Retail + 2nd Location Planning",
    ]
    colours = [RED, GREEN, NAVY, GOLD]
    y = np.arange(len(phases))
    widths = [1, 2, 3, 4]
    starts = [0, 1, 2, 3]
    bars = ax.barh(y, widths, left=starts, color=colours, edgecolor="white", height=0.5)
    for bar, task in zip(bars, tasks):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_y() + bar.get_height()/2,
                task, ha="center", va="center", fontsize=7.5, color="white", fontweight="bold")
    ax.set_yticks(y); ax.set_yticklabels(phases, fontsize=9)
    ax.set_xlim(0, 7.5); ax.set_xticks([])
    ax.set_title("Implementation Roadmap", fontsize=11, fontweight="bold", color=NAVY)
    ax.spines[["top","right","bottom"]].set_visible(False)
    return save_fig("roadmap.png", fig)


print("Generating charts...")
charts = {
    "online_categories": chart_online_categories(),
    "online_impact":     chart_online_impact(),
    "instore_categories":chart_instore_categories(),
    "instore_top3":      chart_instore_top3(),
    "weekly_sales":      chart_weekly_sales(),
    "payment_split":     chart_payment_split(),
    "cost_breakdown":    chart_cost_breakdown(),
    "revenue_target":    chart_revenue_target(),
    "revenue_strategy":  chart_revenue_strategy(),
    "priority_matrix":   chart_priority_matrix(),
    "roadmap":           chart_roadmap(),
}
print(f"  {len(charts)} charts saved to {CHART_DIR}")


# ═════════════════════════════════════════════════════════════════════════════
# 2. WORD DOCUMENT
# ═════════════════════════════════════════════════════════════════════════════

def rgb(h):
    h = h.lstrip("#")
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

NAVY_RGB   = rgb(C_NAVY)
RED_RGB    = rgb(C_RED)
GOLD_RGB   = rgb(C_GOLD)
GREEN_RGB  = rgb(C_GREEN)

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# ── Styles helpers ──
def set_run(run, bold=False, italic=False, size=11, colour=None, font="Calibri"):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    if colour:
        run.font.color.rgb = colour

def heading1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run(run, bold=True, size=18, colour=NAVY_RGB)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), C_RED.lstrip("#"))
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, bold=True, size=13, colour=NAVY_RGB)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(doc, text, bold=False, colour=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(text)
    set_run(run, bold=bold, size=10.5, colour=colour)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(doc, label, text, label_colour=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run(label)
    set_run(r1, bold=True, size=10, colour=label_colour or NAVY_RGB)
    r2 = p.add_run(" — " + text)
    set_run(r2, size=10)
    return p

def callout(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run("⚑  " + label + ":  ")
    set_run(r1, bold=True, size=10, colour=GOLD_RGB)
    r2 = p.add_run(text)
    set_run(r2, size=10)
    # Left border shading via XML
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), C_GOLD.lstrip("#"))
    pBdr.append(left)
    pPr.append(pBdr)
    return p

def insert_chart(doc, path, width_inches=5.5, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = cp.runs[0]
        set_run(run2, italic=True, size=9, colour=rgb(C_GREY))
        cp.paragraph_format.space_after = Pt(10)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run(run, bold=True, size=9.5, colour=RGBColor(255,255,255))
        # Background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), C_NAVY.lstrip("#"))
        tcPr.append(shd)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Data rows
    for ri, row in enumerate(rows):
        drow = table.rows[ri+1]
        fill = "F7F9FC" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = drow.cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run(run, size=9.5)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
    if col_widths:
        for row in table.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)
    doc.add_paragraph()
    return table

def add_agent_table(doc, agents):
    """agents = list of (num, priority_flag, name, desc)"""
    table = doc.add_table(rows=1+len(agents), cols=3)
    table.style = "Table Grid"
    headers = ["#", "Agent Name", "Description"]
    col_widths = [0.4, 1.9, 4.2]
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run(run, bold=True, size=9.5, colour=RGBColor(255,255,255))
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), C_NAVY.lstrip("#"))
        tcPr.append(shd)
    for ri, (num, priority, name, desc) in enumerate(agents):
        drow = table.rows[ri+1]
        fill = "FFF0F3" if priority else ("F7F9FC" if ri%2==0 else "FFFFFF")
        vals = [str(num), ("★ " if priority else "") + name, desc]
        for ci, val in enumerate(vals):
            cell = drow.cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            col = RED_RGB if (ci == 1 and priority) else None
            set_run(run, bold=(ci==1), size=9 if ci==2 else 9.5, colour=col)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
        for ci, w in enumerate(col_widths):
            drow.cells[ci].width = Inches(w)
    for i, w in enumerate(col_widths):
        hrow.cells[i].width = Inches(w)
    doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
print("Building Word document...")

# ── COVER ──
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(20)
r = cover.add_run("AI Agent Strategy Guide")
set_run(r, bold=True, size=28, colour=NAVY_RGB, font="Calibri")

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Mobile Repair Business — Growth & Automation Playbook")
set_run(r2, size=14, colour=RED_RGB, font="Calibri")

doc.add_paragraph()

# Overview stats as a 2×3 table
stat_table = doc.add_table(rows=2, cols=3)
stat_table.style = "Table Grid"
stat_data = [
    ("40 AI Agent Ideas", "20 Online + 20 In-Store"),
    ("£2,314 Avg/week",   "Current Weekly Sales"),
    ("£3,500 Target",     "Weekly Revenue for Expansion"),
    ("£1,818 – £2,834",  "Sales Range (Worst – Best Week)"),
    ("59% Cash / 41% Card", "Payment Split"),
    ("15 Strategies",     "Identified Scaling Best Practices"),
]
fills = [C_NAVY.lstrip("#"), C_RED.lstrip("#"), "10b981", "0ea5e9", "8b5cf6", "f59e0b"]
for ri in range(2):
    for ci in range(3):
        idx = ri*3 + ci
        cell = stat_table.rows[ri].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(stat_data[idx][0] + "\n")
        set_run(r1, bold=True, size=13, colour=RGBColor(255,255,255))
        r2 = p.add_run(stat_data[idx][1])
        set_run(r2, size=9, colour=RGBColor(220,230,255))
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fills[idx])
        tcPr.append(shd)
for row in stat_table.rows:
    for cell in row.cells:
        cell.width = Inches(2.2)

doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════
# PART 1
# ═══════════════════════════════════════════════════
heading1(doc, "PART 1 — 20 AI Agent Ideas: Online / Digital")

# Charts side by side in a 1×2 table
chart_row = doc.add_table(rows=1, cols=2)
chart_row.style = "Table Grid"
for ci, (key, cap) in enumerate([
    ("online_categories", "Online Agents by Category"),
    ("online_impact",     "Revenue Impact vs Ease of Deployment"),
]):
    cell = chart_row.rows[0].cells[ci]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(); run.add_picture(charts[key], width=Inches(3.1))
    cp = cell.add_paragraph(cap)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(cp.runs[0] if cp.runs else cp.add_run(cap), italic=True, size=8, colour=rgb(C_GREY))
doc.add_paragraph()

heading2(doc, "Lead Generation & Traffic  (Agents 1–5)")
add_agent_table(doc, [
    (1, True,  "Repair Cost Estimator Bot",      "Widget where customers describe their problem and get an instant price estimate, capturing their contact info in exchange."),
    (2, False, "Is It Worth Repairing? Advisor",  "Compares repair cost vs. replacement cost for any device, building trust and driving decisions toward booking."),
    (3, False, "Social Media Comment Responder",   "Auto-replies to comments and DMs on Instagram/Facebook with helpful info and a booking link."),
    (4, False, "Local SEO Content Generator",      "Produces weekly blog posts such as 'iPhone screen repair in [City]' to rank on Google searches."),
    (5, False, "Google Review Responder",           "Drafts personalised responses to every review, improving local ranking signals."),
])

heading2(doc, "Booking & Conversion  (Agents 6–10)")
add_agent_table(doc, [
    (6,  True,  "24/7 Booking Chatbot",               "Handles enquiries and books appointments outside business hours, capturing leads that would otherwise be lost."),
    (7,  False, "Abandoned Quote Follow-Up Agent",      "If someone gets a quote but doesn't book, sends a timed follow-up via SMS or email."),
    (8,  False, "Device Triage Agent",                  "Asks diagnostic questions to qualify jobs before they arrive, saving technician time."),
    (9,  False, "Waitlist & Cancellation Filler",       "Automatically fills cancelled slots by texting customers on a waitlist."),
    (10, False, "Warranty Upsell Bot",                  "After booking, pitches protection plans or screen protectors with a one-click add-on."),
])

heading2(doc, "Customer Retention  (Agents 11–15)")
add_agent_table(doc, [
    (11, True,  "Repair Status Updater",           "Sends automated SMS/WhatsApp updates at each stage: received, in repair, ready for collection. Reduces inbound support calls."),
    (12, False, "Post-Repair Check-In Agent",       "Messages customers 7 days after collection to ask how the device is working and request a Google review."),
    (13, False, "Annual Device Health Reminder",    "Reaches out to past customers yearly — 'Your repair was 12 months ago — time for a checkup?'"),
    (14, False, "Loyalty Points Tracker",           "Tracks repeat customers and automatically applies discounts or sends referral codes."),
    (15, False, "Recall & Parts Update Notifier",   "Alerts past customers if a known fault is discovered for their device model."),
])

heading2(doc, "Operations & Upselling  (Agents 16–20)")
add_agent_table(doc, [
    (16, False, "Parts Inventory Assistant",     "Monitors stock levels and flags when common parts are running low, and can auto-order from suppliers."),
    (17, False, "Job Profitability Analyser",    "Logs each repair type, calculates margins, and suggests which jobs to promote or deprioritise."),
    (18, False, "Competitor Price Monitor",      "Scrapes competitor websites weekly and alerts if pricing needs adjusting to stay competitive."),
    (19, False, "Referral Campaign Manager",     "Automatically sends referral links to happy customers and tracks conversions with small rewards."),
    (20, False, "Insurance Claim Helper",        "Guides customers through claiming phone repairs on their home insurance, removing a common barrier to booking."),
])

callout(doc, "Where to start online",
        "Agents 1, 6, and 11 — the Estimator Bot drives traffic, the 24/7 Booking Chatbot converts it, "
        "and the Repair Status Updater reduces support call volume.")

doc.add_page_break()

# ═══════════════════════════════════════════════════
# PART 2
# ═══════════════════════════════════════════════════
heading1(doc, "PART 2 — 20 AI Agent Ideas: Physical Shop Floor")

chart_row2 = doc.add_table(rows=1, cols=2)
chart_row2.style = "Table Grid"
for ci, (key, cap) in enumerate([
    ("instore_categories", "In-Store Agents by Category"),
    ("instore_top3",       "Top 3 In-Store Agents — Priority Scores"),
]):
    cell = chart_row2.rows[0].cells[ci]
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(); run.add_picture(charts[key], width=Inches(3.1))
    cp = cell.add_paragraph(cap); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(cp.runs[0] if cp.runs else cp.add_run(cap), italic=True, size=8, colour=rgb(C_GREY))
doc.add_paragraph()

heading2(doc, "Customer Check-In & Reception  (Agents 1–5)")
add_agent_table(doc, [
    (1, True,  "Smart Check-In Kiosk",        "Tablet at the front desk where customers tap their device model, describe the fault, and receive a repair ticket — no staff needed for intake."),
    (2, False, "Queue Management Agent",        "Displays estimated wait times and texts customers when their repair is next, so they can step out instead of waiting."),
    (3, False, "Walk-In Triage Assistant",      "Runs a quick diagnostic questionnaire and tells the customer their likely repair time and cost before speaking to a technician."),
    (4, False, "Multilingual Reception Bot",    "Greets walk-in customers in their preferred language — especially valuable in diverse urban areas like London."),
    (5, False, "Appointment Check-In Agent",    "Recognises returning customers by phone number, pulls up their history, and alerts the technician they have arrived."),
])

heading2(doc, "In-Store Sales & Upselling  (Agents 6–10)")
add_agent_table(doc, [
    (6,  False, "Accessories Recommender Screen",  "A display showing 'Customers with your phone also bought…' — cases, chargers, screen protectors — while they wait."),
    (7,  False, "Trade-In Value Agent",             "Customers enter old device details into a kiosk and instantly receive a trade-in quote, encouraging a refurbished handset upsell."),
    (8,  False, "Insurance Add-On Pitcher",         "After the repair is logged, automatically presents a 30-second pitch for a protection plan on a tablet."),
    (9,  False, "Bundle Deal Suggester",            "If a customer books a screen repair, the agent flags that a simultaneous battery replacement saves labour cost."),
    (10, False, "Refurbished Phone Browser",        "An in-store touchscreen kiosk showing current refurbished stock with specs, prices, and a reserve button."),
])

heading2(doc, "Technician Support  (Agents 11–15)")
add_agent_table(doc, [
    (11, True,  "AI Repair Guide Assistant",        "Technicians type a fault description and the agent instantly pulls up the best repair walkthrough, known issues, and part numbers."),
    (12, False, "Fault Diagnosis Agent",             "Technician inputs symptoms and the agent cross-references known faults to suggest the most likely cause."),
    (13, False, "Parts Lookup & Ordering Agent",     "Scans a device model and tells the technician which parts are in stock and which need ordering, with one-tap ordering."),
    (14, False, "Repair Time Estimator for Techs",   "Based on job complexity and current workload, tells the technician how long to quote — keeping promises realistic."),
    (15, False, "Quality Control Checklist Agent",   "After a repair, walks the technician through a pass/fail checklist before handing the device back to the customer."),
])

heading2(doc, "Customer Handoff & After-Service  (Agents 16–20)")
add_agent_table(doc, [
    (16, False, "Repair Summary Explainer",      "When handing back the device, a screen or printout explains in plain English what was fixed and provides aftercare tips."),
    (17, True,  "On-the-Spot Review Requester",  "As the customer pays, a tablet prompts them to leave a Google review while they're still feeling positive about getting their phone back."),
    (18, False, "Payment Plan Agent",            "For larger repairs, calculates and presents a simple 'pay in 3' instalment option and sends the agreement to their phone."),
    (19, False, "Data Backup Reminder Kiosk",    "Before repairs begin, alerts the customer if their device has no recent backup and offers a paid backup service on the spot."),
    (20, False, "Uncollected Device Manager",    "Tracks devices not collected after repair and sends escalating reminders (day 3, 7, 14) via SMS, flagging when to apply storage fees."),
])

callout(doc, "Top 3 in-store",
        "Agent #1 Smart Check-In Kiosk (frees up staff), Agent #11 AI Repair Guide (faster technicians), "
        "Agent #17 On-the-Spot Review Requester (Google reviews drive local foot traffic).")

doc.add_page_break()

# ═══════════════════════════════════════════════════
# PART 3
# ═══════════════════════════════════════════════════
heading1(doc, "PART 3 — Financial Analysis: Key Findings")

# Stats summary table
add_table(doc,
    ["Metric", "Value", "Notes"],
    [
        ["Average weekly total sales", "£2,314",         "Range: £1,818 – £2,834"],
        ["Cash vs card split",          "59% / 41%",      "Cash-heavy — track carefully"],
        ["Best week",                   "£2,834",         "Mar 2–8"],
        ["Worst week",                  "£1,818",         "Mar 30–Apr 5"],
        ["Highest-earning days",        "Sat & Sun",      "Consistently highest cash takings"],
        ["Weekly salary cost",          "£800 (fixed)",   "~35% of total weekly revenue"],
    ],
    col_widths=[2.2, 1.5, 2.8]
)

# Charts row
chart_row3 = doc.add_table(rows=1, cols=3)
chart_row3.style = "Table Grid"
for ci, (key, cap) in enumerate([
    ("weekly_sales",   "Weekly Sales Performance"),
    ("payment_split",  "Payment Method Split"),
    ("cost_breakdown", "Weekly Cost Breakdown"),
]):
    cell = chart_row3.rows[0].cells[ci]
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(); run.add_picture(charts[key], width=Inches(2.1))
    cp = cell.add_paragraph(cap); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(cp.runs[0] if cp.runs else cp.add_run(cap), italic=True, size=7.5, colour=rgb(C_GREY))
doc.add_paragraph()

heading2(doc, "Most Common Expenses")
add_table(doc,
    ["Expense", "Amount / Frequency", "Priority Action"],
    [
        ["Salaries",             "£800/week (fixed)",     "Dominant cost — ~35% of revenue; review full-time vs part-time split"],
        ["Bangla Tech",          "Almost every week",     "Most frequent parts supplier — negotiate bulk / 30-day terms"],
        ["Revive",               "Regular",               "Second most frequent supplier — bundle orders with Bangla Tech negotiations"],
        ["eBay",                 "Regular",               "Used for phones, screens, batteries — consider a dedicated seller account"],
        ["HMRC Tax",             "Periodic — £255",       "Set aside weekly; do not allow it to create a cash shock"],
        ["Council Rates",        "Periodic — £80–£178",   "Fixed overhead — budget monthly"],
        ["Insurance / Ria",      "Periodic",              "Track separately to understand true operational cost"],
    ],
    col_widths=[1.6, 1.8, 3.1]
)

heading2(doc, "Financial Concerns")
callout(doc, "Cash Flow Risk",
        "End-of-week cash is often dangerously thin — several weeks show negative cash balances before personal "
        "top-ups. A minimum cash reserve of £1,600 (2 weeks of salary) is essential to prevent operational disruption.")
callout(doc, "External Cash Dependency",
        "The business is heavily reliant on personal cash injections to cover high-expense salary weeks. "
        "Separating business and personal finances is a critical first step.")
callout(doc, "Side Venture Not Separated",
        "Perfume stock appears as a side venture within the same books. Track it separately to accurately measure "
        "core repair business profitability and avoid distorted P&L figures.")

insert_chart(doc, charts["revenue_target"], width_inches=5.0,
             caption="Current Weekly Revenue vs Targets — £2,314 current → £3,500 to unlock second location")

doc.add_page_break()

# ═══════════════════════════════════════════════════
# PART 4
# ═══════════════════════════════════════════════════
heading1(doc, "PART 4 — Scaling Strategy: 15 Best Practices")

insert_chart(doc, charts["revenue_strategy"], width_inches=6.0,
             caption="Estimated Additional Monthly Revenue by Scaling Strategy (Min and Max)")

insert_chart(doc, charts["priority_matrix"], width_inches=5.5,
             caption="Strategy Priority Matrix — plot of Revenue Impact vs Effort Required")

strategies = [
    ("1",  "Fix the Cash Flow Crisis First",
     "Separate business and personal finances immediately. Build a cash reserve equal to 2 weeks of salary (£1,600). "
     "Rule: if end-of-week cash is below £200, no discretionary spending that week."),
    ("2",  "Negotiate Bulk Supplier Pricing",
     "Approach Bangla Tech and Revive for monthly accounts with 30-day payment terms and volume discounts. "
     "Place one bulk order per month for the top 10 most-used parts."),
    ("3",  "Increase Average Order Value",
     "Train staff to offer battery + screen repairs together. Target an average transaction value of £80+ "
     "versus the current estimated £40–60."),
    ("4",  "Launch a WhatsApp Broadcast List",
     "Message 200+ past customers instantly at zero cost. One broadcast per week (tip, offer, or reminder) "
     "keeps the shop top of mind and drives repeat visits."),
    ("5",  "Formalise the eBay Sales Operation",
     "Set up a proper eBay store for refurbished phones, spare parts, and accessories. "
     "Even 5–10 eBay sales per week at £20–50 margin = £400–1,000/month additional revenue."),
    ("6",  "Introduce a Monthly Repair Membership",
     "£6–8/month provides: priority service, one free screen protector fitting, and 10% off all repairs. "
     "50 members = £300/month in guaranteed recurring income."),
    ("7",  "Launch a Refurbished Phone Retail Line",
     "Formalise phone-flipping with a dedicated float of £500–1,000 in stock. Buy at £40–80, sell at £100–150. "
     "Potential: £400–800/month additional margin."),
    ("8",  "Extend Trading Hours on Weekends",
     "Saturdays and Sundays consistently show the highest takings. Extending to 7pm could capture £50–100 "
     "extra per day — £400–800/month additional revenue."),
    ("9",  "Target the Bengali/South Asian Community",
     "Market in Bengali via WhatsApp broadcast lists, local Facebook groups, and leaflets in local shops. "
     "Community trust translates into referrals at near-zero marketing cost."),
    ("10", "Hire Part-Time Instead of Full-Time",
     "If one person is full-time and one is weekend-only, the weekly salary bill drops from £800 to £550–600 "
     "— freeing £800–1,000/month for reinvestment into the business."),
    ("11", "Register on eBay as a Seller",
     "Extend reach nationally with a dedicated eBay store for refurbished phones, spare parts, and accessories. "
     "Complements Strategy 5 with a formal, branded storefront."),
    ("12", "Introduce a Simple Loyalty Card",
     "A stamp card (10 repairs = one free screen protector fitting) costs almost nothing to print but "
     "significantly increases return visit rates and customer lifetime value."),
    ("13", "Build a Google Business Profile",
     "A fully optimised profile with 20–30 genuine reviews is the single highest-ROI marketing action available. "
     "Each 5-star review drives an estimated 1–3 additional customers per month."),
    ("14", "Separate Business Finances Completely",
     "Use a dedicated business bank account and a bookkeeping app (QuickBooks Self-Employed at £8/month, "
     "or free Wave). Enables proper P&L visibility and makes tax returns significantly faster."),
    ("15", "Plan for a Second Location After £3,500/week",
     "The current average of £2,314/week needs to reach £3,500 consistently before a second location is "
     "financially viable. Achievable through the combined strategies above."),
]
add_table(doc,
    ["#", "Strategy", "Detail"],
    [[s[0], s[1], s[2]] for s in strategies],
    col_widths=[0.3, 1.9, 4.3]
)

doc.add_paragraph()
insert_chart(doc, charts["roadmap"], width_inches=6.5,
             caption="Implementation Roadmap — Four-phase plan from Month 1 to Month 6+")

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(DOCX_PATH)
print(f"Word document saved: {DOCX_PATH}")

# ═════════════════════════════════════════════════════════════════════════════
# 3. PDF via LibreOffice
# ═════════════════════════════════════════════════════════════════════════════
print("Converting to PDF via LibreOffice...")
result = subprocess.run(
    ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", OUT_DIR, DOCX_PATH],
    capture_output=True, text=True, timeout=120
)
if result.returncode != 0:
    print("LibreOffice stderr:", result.stderr)
    raise RuntimeError("PDF conversion failed")

PDF_PATH = DOCX_PATH.replace(".docx", ".pdf")
if os.path.exists(PDF_PATH):
    print(f"PDF saved:          {PDF_PATH}")
else:
    print("Warning: PDF not found at expected path")
    print(result.stdout)

print("\nDone.")
print(f"  DOCX → {DOCX_PATH}")
print(f"  PDF  → {PDF_PATH}")
